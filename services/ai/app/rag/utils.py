"""
Document processing utilities for RAG system.

Provides functions for extracting text from various document formats (PDF, DOCX, TXT, MD)
and chunking text into smaller pieces suitable for embedding and retrieval.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    Image = None
    convert_from_path = None


def extract_text_from_pdf(file_path: str, use_ocr: bool = True) -> str:
    """
    Extract text content from a PDF file.
    First attempts regular text extraction, then falls back to OCR if needed.
    
    Args:
        file_path: Path to the PDF file
        use_ocr: If True, use OCR when regular extraction yields minimal text
        
    Returns:
        Extracted text as a string, or empty string if extraction fails
    """
    if pdfplumber is None:
        raise ImportError("pdfplumber is required for PDF extraction. Install with: pip install pdfplumber")
    
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        extracted_text = "\n".join(text_parts).strip()
        
        # If regular extraction yields minimal text and OCR is available, try OCR
        if use_ocr and OCR_AVAILABLE and len(extracted_text) < 500:
            print(f"PDF {file_path} yielded minimal text ({len(extracted_text)} chars), attempting OCR...")
            try:
                ocr_text = extract_text_from_pdf_ocr(file_path)
                if len(ocr_text) > len(extracted_text):
                    print(f"OCR extracted {len(ocr_text)} characters (vs {len(extracted_text)} from regular extraction)")
                    return ocr_text
            except Exception as ocr_e:
                print(f"OCR failed for PDF {file_path}: {ocr_e}, using regular extraction")
        
        return extracted_text
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        # Try OCR as fallback if regular extraction completely fails
        if use_ocr and OCR_AVAILABLE:
            try:
                print(f"Attempting OCR fallback for PDF {file_path}...")
                return extract_text_from_pdf_ocr(file_path)
            except Exception as ocr_e:
                print(f"OCR fallback also failed: {ocr_e}")
        return ""


def extract_text_from_pdf_ocr(file_path: str) -> str:
    """
    Extract text from PDF using OCR (Optical Character Recognition).
    Converts PDF pages to images and uses Tesseract OCR.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    if not OCR_AVAILABLE:
        raise ImportError("OCR libraries required. Install with: pip install pytesseract pdf2image pillow")
    
    try:
        print(f"Converting PDF {file_path} to images for OCR...")
        # Convert PDF pages to images
        images = convert_from_path(file_path, dpi=300)  # Higher DPI for better OCR quality
        
        text_parts = []
        total_pages = len(images)
        print(f"Processing {total_pages} pages with OCR...")
        
        for i, image in enumerate(images, 1):
            if i % 10 == 0 or i == 1:
                print(f"  OCR processing page {i}/{total_pages}...")
            # Use Tesseract OCR
            page_text = pytesseract.image_to_string(image, lang='eng')
            if page_text.strip():
                text_parts.append(f"--- Page {i} ---\n{page_text.strip()}")
        
        result = "\n\n".join(text_parts).strip()
        print(f"OCR extraction complete: {len(result)} characters from {total_pages} pages")
        return result
    except Exception as e:
        print(f"Error in OCR extraction for PDF {file_path}: {e}")
        raise


def extract_text_from_docx(file_path: str, use_ocr: bool = True) -> str:
    """
    Extract text content from a DOCX file.
    First attempts regular text extraction, then falls back to OCR if needed.
    
    Args:
        file_path: Path to the DOCX file
        use_ocr: If True, use OCR when regular extraction yields minimal text
        
    Returns:
        Extracted text as a string, or empty string if extraction fails
    """
    if Document is None:
        raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
    
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        extracted_text = "\n".join(text_parts).strip()
        
        # If regular extraction yields minimal text and OCR is available, try OCR
        if use_ocr and OCR_AVAILABLE and len(extracted_text) < 500:
            print(f"DOCX {file_path} yielded minimal text ({len(extracted_text)} chars), attempting OCR...")
            try:
                ocr_text = extract_text_from_docx_ocr(file_path)
                if len(ocr_text) > len(extracted_text):
                    print(f"OCR extracted {len(ocr_text)} characters (vs {len(extracted_text)} from regular extraction)")
                    return ocr_text
            except Exception as ocr_e:
                print(f"OCR failed for DOCX {file_path}: {ocr_e}, using regular extraction")
        
        return extracted_text
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        # Try OCR as fallback if regular extraction completely fails
        if use_ocr and OCR_AVAILABLE:
            try:
                print(f"Attempting OCR fallback for DOCX {file_path}...")
                return extract_text_from_docx_ocr(file_path)
            except Exception as ocr_e:
                print(f"OCR fallback also failed: {ocr_e}")
        return ""


def extract_text_from_docx_ocr(file_path: str) -> str:
    """
    Extract text from DOCX using OCR (Optical Character Recognition).
    Converts DOCX to PDF first, then to images, and uses Tesseract OCR.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
    """
    if not OCR_AVAILABLE:
        raise ImportError("OCR libraries required. Install with: pip install pytesseract pdf2image pillow")
    
    try:
        # For DOCX, we need to convert to PDF first, then to images
        # Using python-docx2pdf or libreoffice, but simpler: convert to images directly
        # Alternative: use docx2pdf library, but for now we'll use a workaround
        
        # Method: Convert DOCX to images by rendering each page
        # This requires additional libraries, so we'll use a simpler approach:
        # Try to extract images from DOCX and OCR them, plus any embedded text
        
        print(f"Extracting images from DOCX {file_path} for OCR...")
        doc = Document(file_path)
        text_parts = []
        
        # First, get any text that was extracted
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract images from the document and OCR them
        # Note: python-docx doesn't directly support image extraction
        # We'll need to use a different approach
        
        # Alternative: Convert DOCX to PDF using libreoffice or docx2pdf
        # Then use PDF OCR method
        import tempfile
        import subprocess
        import os
        
        # Try to convert DOCX to PDF using libreoffice (if available)
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
                tmp_pdf_path = tmp_pdf.name
            
            # Use libreoffice to convert DOCX to PDF
            # LibreOffice converts to the same directory as the input file
            # So we need to convert in the temp directory
            import shutil
            temp_dir = os.path.dirname(tmp_pdf_path)
            temp_docx = os.path.join(temp_dir, os.path.basename(file_path))
            shutil.copy2(file_path, temp_docx)
            
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, temp_docx],
                capture_output=True,
                timeout=300,
                env={**os.environ, 'HOME': temp_dir}  # Set HOME to avoid user config issues
            )
            
            # Clean up temp DOCX
            if os.path.exists(temp_docx):
                os.unlink(temp_docx)
            
            # Find the generated PDF (LibreOffice uses the same name with .pdf extension)
            expected_pdf = os.path.join(temp_dir, os.path.splitext(os.path.basename(file_path))[0] + '.pdf')
            if os.path.exists(expected_pdf):
                tmp_pdf_path = expected_pdf
            
            if result.returncode == 0 and os.path.exists(tmp_pdf_path):
                print(f"Converted DOCX to PDF, now using OCR on PDF...")
                pdf_text = extract_text_from_pdf_ocr(tmp_pdf_path)
                os.unlink(tmp_pdf_path)  # Clean up temp file
                return pdf_text
            else:
                print(f"libreoffice conversion failed, trying alternative method...")
                if os.path.exists(tmp_pdf_path):
                    os.unlink(tmp_pdf_path)
        except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as conv_e:
            print(f"DOCX to PDF conversion failed: {conv_e}")
            # Fallback: return whatever text we extracted
            if text_parts:
                return "\n".join(text_parts).strip()
            raise Exception(f"Could not convert DOCX to PDF for OCR: {conv_e}")
        
        # If we get here, return extracted text
        return "\n".join(text_parts).strip() if text_parts else ""
        
    except Exception as e:
        print(f"Error in OCR extraction for DOCX {file_path}: {e}")
        raise


def chunk_text(
    text: str,
    chunk_size: int = 900,
    overlap: int = 120,
    separator: str = "\n\n"
) -> Iterable[str]:
    """
    Split text into overlapping chunks for embedding.
    
    This function intelligently chunks text by:
    1. First attempting to split on paragraph boundaries (double newlines)
    2. If chunks are still too large, splitting on sentences
    3. If still too large, splitting on words
    4. As a last resort, splitting on characters
    
    Args:
        text: The text to chunk
        chunk_size: Target size for each chunk (in characters)
        overlap: Number of characters to overlap between chunks
        separator: Preferred separator for splitting (default: paragraph breaks)
        
    Yields:
        Text chunks as strings
    """
    if not text or not text.strip():
        return
    
    text = text.strip()
    
    # If text is smaller than chunk_size, return as-is
    if len(text) <= chunk_size:
        yield text
        return
    
    # Strategy 1: Split by paragraphs (double newlines)
    paragraphs = text.split(separator)
    current_chunk = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # If adding this paragraph would exceed chunk_size
        if current_chunk and len(current_chunk) + len(separator) + len(para) > chunk_size:
            # Yield current chunk
            if current_chunk:
                yield current_chunk
            
            # Start new chunk with overlap
            if overlap > 0 and len(current_chunk) > overlap:
                # Take last 'overlap' characters from previous chunk
                overlap_text = current_chunk[-overlap:]
                current_chunk = overlap_text + separator + para
            else:
                current_chunk = para
        else:
            # Add paragraph to current chunk
            if current_chunk:
                current_chunk += separator + para
            else:
                current_chunk = para
    
    # Yield remaining chunk
    if current_chunk:
        yield current_chunk
    
    # If we still have very long chunks, split them further
    # This handles cases where individual paragraphs are longer than chunk_size
    for chunk in _split_long_chunks(text, chunk_size, overlap):
        if chunk not in (current_chunk,):  # Avoid duplicates
            yield chunk


def _split_long_chunks(text: str, chunk_size: int, overlap: int) -> Iterable[str]:
    """
    Split text that's longer than chunk_size into smaller pieces.
    
    This is a fallback for when paragraph-based chunking doesn't work.
    """
    if len(text) <= chunk_size:
        yield text
        return
    
    # Split by sentences first
    import re
    sentences = re.split(r'([.!?]+\s+)', text)
    
    # Recombine sentences with their punctuation
    combined_sentences = []
    for i in range(0, len(sentences) - 1, 2):
        if i + 1 < len(sentences):
            combined_sentences.append(sentences[i] + sentences[i + 1])
        else:
            combined_sentences.append(sentences[i])
    
    current_chunk = ""
    for sentence in combined_sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        
        if current_chunk and len(current_chunk) + len(sentence) + 1 > chunk_size:
            if current_chunk:
                yield current_chunk
            # Start new chunk with overlap
            if overlap > 0 and len(current_chunk) > overlap:
                current_chunk = current_chunk[-overlap:] + " " + sentence
            else:
                current_chunk = sentence
        else:
            current_chunk = (current_chunk + " " + sentence).strip() if current_chunk else sentence
    
    if current_chunk:
        yield current_chunk
    
    # Final fallback: character-based splitting if still needed
    if len(text) > chunk_size * 2:
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if chunk.strip():
                yield chunk.strip()
            start = end - overlap if overlap > 0 else end

