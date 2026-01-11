"""
Document loader module for Zamfara RAG system.

Handles recursive loading of documents from the Zamfara directory,
supporting PDF, DOCX, TXT, and Markdown formats with OCR fallback.
"""

from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Generator, Optional, Any

# PDF extraction
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX extraction
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# OCR support
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    Image = None
    convert_from_path = None

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Represents a loaded document with extracted content and metadata."""
    
    # Core content
    text: str
    
    # File metadata
    file_path: Path
    file_name: str
    file_extension: str
    file_size_bytes: int
    
    # Temporal metadata
    last_modified: datetime
    created_at: Optional[datetime] = None
    
    # Extraction metadata
    extraction_method: str = "standard"  # standard, ocr, hybrid
    page_count: int = 0
    word_count: int = 0
    
    # Inferred metadata (populated later)
    document_title: str = ""
    document_type: str = "unknown"
    department: str = "unknown"
    
    def __post_init__(self):
        """Calculate word count if not set."""
        if self.word_count == 0 and self.text:
            self.word_count = len(self.text.split())
        if not self.document_title:
            # Extract title from filename
            self.document_title = self.file_name.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")


class DocumentLoader:
    """
    Loads documents from the Zamfara directory with support for multiple formats.
    
    Features:
    - Recursive directory scanning
    - Multi-format support (PDF, DOCX, TXT, MD)
    - OCR fallback for scanned documents
    - Metadata extraction
    """
    
    def __init__(
        self,
        source_dir: Path,
        supported_formats: set[str] = None,
        max_file_size_mb: int = 50,
        use_ocr: bool = True
    ):
        """
        Initialize the document loader.
        
        Args:
            source_dir: Root directory to load documents from
            supported_formats: Set of supported file extensions
            max_file_size_mb: Maximum file size to process
            use_ocr: Whether to use OCR for scanned documents
        """
        self.source_dir = Path(source_dir)
        self.supported_formats = supported_formats or {".pdf", ".docx", ".txt", ".md"}
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.use_ocr = use_ocr and OCR_AVAILABLE
        
        if not self.source_dir.exists():
            logger.warning(f"Source directory does not exist: {self.source_dir}")
    
    def load_all(self) -> Generator[LoadedDocument, None, None]:
        """
        Recursively load all documents from the source directory.
        
        Yields:
            LoadedDocument objects for each successfully loaded file
        """
        if not self.source_dir.exists():
            logger.error(f"Source directory does not exist: {self.source_dir}")
            return
        
        logger.info(f"Loading documents from: {self.source_dir}")
        
        for file_path in self._iter_files():
            try:
                doc = self.load_file(file_path)
                if doc and doc.text.strip():
                    logger.info(f"Loaded: {file_path.name} ({doc.word_count} words)")
                    yield doc
                else:
                    logger.warning(f"Empty or failed to load: {file_path.name}")
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
                continue
    
    def _iter_files(self) -> Generator[Path, None, None]:
        """Iterate over all supported files in the source directory."""
        for path in self.source_dir.rglob("*"):
            if not path.is_file():
                continue
            if path.suffix.lower() not in self.supported_formats:
                continue
            if path.stat().st_size > self.max_file_size_bytes:
                logger.warning(f"Skipping large file: {path.name}")
                continue
            yield path
    
    def load_file(self, file_path: Path) -> Optional[LoadedDocument]:
        """
        Load a single document file.
        
        Args:
            file_path: Path to the file to load
            
        Returns:
            LoadedDocument or None if loading fails
        """
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        stat = file_path.stat()
        suffix = file_path.suffix.lower()
        
        # Extract text based on file type
        text = ""
        extraction_method = "standard"
        page_count = 0
        
        if suffix == ".pdf":
            text, page_count, extraction_method = self._extract_pdf(file_path)
        elif suffix == ".docx":
            text, page_count, extraction_method = self._extract_docx(file_path)
        elif suffix in {".txt", ".md"}:
            text = self._extract_text_file(file_path)
            page_count = 1
        else:
            logger.warning(f"Unsupported format: {suffix}")
            return None
        
        return LoadedDocument(
            text=text,
            file_path=file_path,
            file_name=file_path.name,
            file_extension=suffix,
            file_size_bytes=stat.st_size,
            last_modified=datetime.fromtimestamp(stat.st_mtime),
            extraction_method=extraction_method,
            page_count=page_count,
        )
    
    def _extract_pdf(self, file_path: Path) -> tuple[str, int, str]:
        """
        Extract text from PDF, with OCR fallback.
        
        Returns:
            Tuple of (text, page_count, extraction_method)
        """
        if pdfplumber is None:
            raise ImportError("pdfplumber is required for PDF extraction")
        
        text_parts = []
        page_count = 0
        
        try:
            with pdfplumber.open(str(file_path)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            extracted_text = "\n\n".join(text_parts).strip()
            
            # If minimal text extracted and OCR available, try OCR
            if self.use_ocr and len(extracted_text) < 500:
                logger.info(f"Attempting OCR for: {file_path.name}")
                try:
                    ocr_text = self._extract_pdf_ocr(file_path)
                    if len(ocr_text) > len(extracted_text):
                        return ocr_text, page_count, "ocr"
                except Exception as e:
                    logger.warning(f"OCR failed for {file_path.name}: {e}")
            
            return extracted_text, page_count, "standard"
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            
            # Try OCR as complete fallback
            if self.use_ocr:
                try:
                    ocr_text = self._extract_pdf_ocr(file_path)
                    return ocr_text, page_count, "ocr"
                except Exception:
                    pass
            
            return "", 0, "failed"
    
    def _extract_pdf_ocr(self, file_path: Path) -> str:
        """Extract text from PDF using OCR."""
        if not OCR_AVAILABLE:
            raise ImportError("OCR libraries not available")
        
        images = convert_from_path(str(file_path), dpi=300)
        text_parts = []
        
        for i, image in enumerate(images, 1):
            page_text = pytesseract.image_to_string(image, lang='eng')
            if page_text.strip():
                text_parts.append(page_text.strip())
        
        return "\n\n".join(text_parts).strip()
    
    def _extract_docx(self, file_path: Path) -> tuple[str, int, str]:
        """
        Extract text from DOCX file.
        
        Returns:
            Tuple of (text, page_count, extraction_method)
        """
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX extraction")
        
        try:
            doc = DocxDocument(str(file_path))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts).strip()
            
            # Estimate page count (rough approximation)
            page_count = max(1, len(text) // 3000)
            
            return text, page_count, "standard"
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return "", 0, "failed"
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text or markdown files."""
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception as e:
            logger.error(f"Text file extraction failed for {file_path}: {e}")
            return ""
    
    def get_document_count(self) -> int:
        """Count the number of supported documents in the source directory."""
        return sum(1 for _ in self._iter_files())




